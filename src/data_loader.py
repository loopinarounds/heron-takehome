import os
import PyPDF2
from PIL import Image
import pytesseract
from docx import Document

def extract_text_from_pdf(file_path_or_file):
    if hasattr(file_path_or_file, 'read'):
        reader = PyPDF2.PdfReader(file_path_or_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    else:
        with open(file_path_or_file, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
        return text

def extract_text_from_docx(file_path_or_file):
    if hasattr(file_path_or_file, 'read'):
        doc = Document(file_path_or_file)
    else:
        doc = Document(file_path_or_file)
    
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_image(file_path_or_file):
    if hasattr(file_path_or_file, 'read'):
        image = Image.open(file_path_or_file)
    else:

        image = Image.open(file_path_or_file)
    
    text = pytesseract.image_to_string(image)
    return text

def load_data(data_dir):
    texts = []
    labels = []
    for file_name in os.listdir(data_dir):
        file_path = os.path.join(data_dir, file_name)
        
        if file_name.endswith('.pdf'):
            label = determine_label(file_name) 
            text = extract_text_from_pdf(file_path)
            texts.append(text)
            labels.append(label)  
        elif file_name.endswith('.docx'):
            label = determine_label(file_name)  
            text = extract_text_from_docx(file_path)
            texts.append(text)
            labels.append(label)  
        elif file_name.endswith('.jpg') or file_name.endswith('.jpeg'):
            label = determine_label(file_name) 
            text = extract_text_from_image(file_path)
            texts.append(text)
            labels.append(label) 
        else:
            print(f"Unsupported file type: {file_name}")
            continue 

    print(f"Total texts: {len(texts)}, Total labels: {len(labels)}")
    for i in range(len(texts)):
        print(f"Text {i}: {texts[i][:30]}... | Label: {labels[i]}")  

    return texts, labels

def determine_label(file_name):
    if "drivers_license" in file_name:
        return "drivers_license"
    elif "bank_statement" in file_name:
        return "bank_statement"
    elif "invoice" in file_name:
        return "invoice"
    else:
        return "unknown"